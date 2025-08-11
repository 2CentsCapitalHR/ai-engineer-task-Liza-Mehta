import gradio as gr
from docx import Document
import json
from pathlib import Path

def analyze_docs(files):
    comments = []
    reviewed_files = []

    for file_path in files:  # file_path is a string path
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        if not any("ADGM" in p for p in paragraphs):
            comments.append({
                "document": Path(file_path).name,
                "context": "Jurisdiction",
                "issue": "Jurisdiction not specified as ADGM",
                "severity": "High",
                "suggestion": "Replace jurisdiction clause with 'Abu Dhabi Global Market (ADGM) Courts'."
            })

        if not any("[Company Name]" in p or "{{company_name}}" in p for p in paragraphs):
            comments.append({
                "document": Path(file_path).name,
                "context": "Company Details",
                "issue": "Company name placeholder missing",
                "severity": "Medium",
                "suggestion": "Add '[Company Name]' placeholder where applicable."
            })

        if not any("Registered Address" in p or "registered office" in p.lower() for p in paragraphs):
            comments.append({
                "document": Path(file_path).name,
                "context": "Address",
                "issue": "Registered address missing",
                "severity": "High",
                "suggestion": "Include the company's registered office address."
            })

        if not any("Signature" in p or "Signed by" in p for p in paragraphs):
            comments.append({
                "document": Path(file_path).name,
                "context": "Signatures",
                "issue": "Signature section missing",
                "severity": "High",
                "suggestion": "Add signature lines for authorized signatories."
            })

        if not any("[Date]" in p or "{{date}}" in p for p in paragraphs):
            comments.append({
                "document": Path(file_path).name,
                "context": "Date",
                "issue": "Date placeholder missing",
                "severity": "Low",
                "suggestion": "Add '[Date]' placeholder to indicate document date."
            })

        doc.add_page_break()
        para = doc.add_paragraph("Reviewer Comments")
        para.runs[0].bold = True

        for c in [c for c in comments if c["document"] == Path(file_path).name]:
            doc.add_paragraph(f"- Issue: {c['issue']} (Severity: {c['severity']})")
            doc.add_paragraph(f"  Suggestion: {c['suggestion']}")

        reviewed_path = f"reviewed_{Path(file_path).name}"
        doc.save(reviewed_path)
        reviewed_files.append(reviewed_path)

    json_path = "report.json"
    with open(json_path, "w") as f:
        json.dump({
            "process": "Company Incorporation",
            "documents_uploaded": len(files),
            "issues_found": comments
        }, f, indent=2)

    return reviewed_files, json_path

demo = gr.Interface(
    fn=analyze_docs,
    inputs=gr.File(file_types=[".docx"], file_count="multiple", type="filepath", label="Upload .docx files"),
    outputs=[gr.File(label="Reviewed DOCX", type="filepath"), gr.File(label="JSON Report", type="filepath")],
    title="ADGM Corporate Agent - Demo"
)

if __name__ == "__main__":
    demo.launch()
